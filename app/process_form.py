import re
from docx import Document
import os
import requests


def docx_replace_regex(doc_obj, regex , replace):

    #regex = "/^" + regex + "$/"
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            counter = 0
            for i in range(len(inline)):

                counter=counter+1
                #print inline[i].text
                if regex.search(inline[i].text):
                    #print "found: ",regex,replace,inline[i].text 
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
                    #print replace, regex, inline[i].text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

def send_email(filename,email,subject,text,html):

	print filename,email  
	print requests.post("https://api.mailgun.net/v3/jaskapital.com/messages",  
                    auth=("api", "key-7028ed900aeadbbe928022178c1eb33d"),
                    files=[("attachment", open(filename))],
                    data={"from": "Panitia Kaderisasi <kresna.jenie@gmail.com>",
                          "to": email,
                          "cc": "kresna.jenie@gmail.com",
                          "subject": subject,
                          "text": text,
                          "html": html})

def process_data(request):

	APP_ROOT = os.path.dirname(os.path.abspath(__file__))   # refers to application_top
	APP_STATIC = os.path.join(APP_ROOT, 'static')
	filename = os.path.join(APP_STATIC, "kader.docx")
	doc = Document(filename)

	for v in request.form:
		#v2 = "/^" +v + "$/"
		if v == 'tulang':
			print request.form[v]
		v2 = v
		regex1 = re.compile(v2)
		replace1 = request.form[v]
		print v, request.form[v]
		docx_replace_regex(doc, regex1 , replace1)



	nama_depan = request.form['nama_depan_saya']
	nama_belakang = request.form['nama_belakang_saya']
	filename = "pendaftaran-"+nama_depan+"-"+nama_belakang+".docx"
	fileoutput = os.path.join(APP_STATIC, filename)
	email = request.form['email']
	doc.save(fileoutput)
	subject='Pendaftaran Kaderisasi & Internalisasi OSIS 2018/2019'
	text ='Terima kasih atas pendaftaran kamu, tetapi belum selesai. Diharapkan kepada siswa yang sudah mendaftar untuk print file yang sudah diberikan. File yang sudah diberi belum lengkap karena masih membutuhkan testimoni dan surat pernyataan orang tua serta pas foto. Data yang sudah lengkap harap diberikan ke Panitia Kaderisasi di Ruang Osis dari tanggal 21 November 2017 sampai dengan tanggal 24 November 2017 pukul 14.00 - 16.00. Selama ulangan umum pendaftaran akan ditutup dan akan dibuka kembali pada tanggal 4 November 2017 sampai dengan 7 November 2017.Jika ada gangguan harap menghubungi Kresna Jenie, Ignatius Bima, dan Aswandatu Putra secepatnya.'
	html='<html>Terima kasih atas pendaftaran kamu, tetapi belum selesai. Diharapkan kepada siswa yang sudah mendaftar untuk print file yang sudah diberikan. File yang sudah diberi <b><u> belum lengkap </u></b> karena masih membutuhkan testimoni dan surat pernyataan orang tua serta pas foto. <br><br> Data yang sudah lengkap harap diberikan ke Panitia Kaderisasi di Ruang Osis dari tanggal 21 November 2017 sampai dengan tanggal 24 November 2017 pukul 14.00 - 16.00. Selama ulangan umum pendaftaran akan ditutup dan akan dibuka kembali pada tanggal 4 November 2017 sampai dengan 7 November 2017.<br><br> Jika ada gangguan harap menghubungi <a href="mailto:kresna.jenie@gmail.com">Kresna Jenie</a>, Ignatius Bima, dan Aswandatu Putra secepatnya. </h4></html>'
	send_email(fileoutput,email,subject,text,html)
    




